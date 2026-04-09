"""Word-Dokumente (.docx) einlesen."""

from pathlib import Path
from docx import Document


def parse_word(file_path: Path) -> list[dict]:
    """Word-Datei einlesen. Gibt Absaetze und Tabellen zurueck."""
    doc = Document(str(file_path))
    results = []

    # Absaetze
    current_section = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Ueberschriften als Sections tracken
        if para.style.name.startswith("Heading"):
            current_section = text

        results.append({
            "text": text,
            "content_type": "text",
            "section": current_section,
            "style": para.style.name,
        })

    # Tabellen
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        if rows:
            # Als Text formatieren
            header = " | ".join(rows[0]) if rows else ""
            body = "\n".join(" | ".join(row) for row in rows[1:])
            table_text = f"{header}\n{'─' * len(header)}\n{body}" if body else header

            results.append({
                "text": table_text,
                "content_type": "table",
                "section": f"Tabelle {table_idx + 1}",
                "table_data": rows,
            })

    return results
